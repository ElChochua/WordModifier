from docx import Document
from pypdf import PdfWriter, PdfReader
import os
import subprocess
import shutil
import pandas as pd
from docx.oxml.ns import qn

"""
    Generador de Cartas por Asesor

    Uso:
        python generar_cartas.py

    Requisitos:
        pip install python-docx pypdf
        LibreOffice instalado (para la conversión a PDF)

        Windows:  https://www.libreoffice.org/download/download/
        Mac:      brew install --cask libreoffice
        Linux:    sudo apt install libreoffice
"""
class FileManager:
    def __init__(self, template, output_folder=None, combined_pdf=False, markers=None):
        self.template = template
        self.output_folder = output_folder
        self.combined_pdf = combined_pdf
        self.markers = markers


    def remove_protection(self, doc):
        settings = doc.settings.element
        for tag in ("w:documentProtection", "w:writeProtection"):
            element = settings.find(qn(tag))
            if element is not None:
                settings.remove(element)

    def unlock_document(self, doc_path, output_folder=None):

        doc = Document(doc_path)
        self.remove_protection(doc)

        if output_folder is None:
            output_path = doc_path
        elif os.path.isdir(output_folder):
            output_path = os.path.join(output_folder, os.path.basename(doc_path))
        else:
            output_path = output_folder

        doc.save(output_path)
        return output_path

    def unlock_folder(self, folder_path):
        unlocked = []
        for file_name in os.listdir(folder_path):
            if file_name.lower().endswith(".docx"):
                path = os.path.join(folder_path, file_name)
                self.unlock_document(path)
                unlocked.append(path)
        return unlocked

    def replace_on_paragraph(self,paragraph, replacement):
        full_text = "".join(run.text for run in paragraph.runs)
        if not any(m in full_text for m in replacement):
            return
        new_text = full_text
        for marker, value in replacement.items():
            new_text = new_text.replace(marker, value)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""


    def replace_in_doc(self, doc, replacement):
        for paragraph in doc.paragraphs:
            self.replace_on_paragraph(paragraph, replacement)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self.replace_on_paragraph(paragraph, replacement)
        for section in doc.sections:
            for container in [
                section.header, section.first_page_header, section.even_page_header,
                section.footer, section.first_page_footer, section.even_page_footer,
            ]:
                if container is None:
                    continue
                for paragraph in container.paragraphs:
                    self.replace_on_paragraph(paragraph, replacement)
                for table in container.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                self.replace_on_paragraph(paragraph, replacement)


    def generate_doc(self,template_path, name):
        doc = Document(template_path)

        file_name = name.replace(" ", "_").replace("/", "-")
        output_path = os.path.join(self.output_folder, f"Carta_{file_name}.docx")
        doc.save(output_path)
        return output_path

    def _safe_file_name(self, value):
        file_name = str(value).strip()
        if not file_name:
            file_name = "documento"
        return file_name.replace(" ", "_").replace("/", "-").replace("\\", "-")

    def _row_replacement(self, row):
        if self.markers:
            columns = [marker for marker in self.markers if marker in row.index]
        else:
            columns = list(row.index)

        replacement = {}
        for column in columns:
            value = row[column]
            replacement[str(column)] = "" if pd.isna(value) else str(value)
        return replacement

    def generate_documents(self, data, filename_column=None):
        if self.output_folder is None:
            raise ValueError("output_folder is required to generate documents")

        if not isinstance(data, pd.DataFrame):
            raise TypeError("data must be a pandas DataFrame")

        if data.empty:
            return []

        os.makedirs(self.output_folder, exist_ok=True)

        generated_files = []
        for row_index, (_, row) in enumerate(data.iterrows()):
            replacement = self._row_replacement(row)
            doc = Document(self.template)
            self.replace_in_doc(doc, replacement)

            if filename_column and filename_column in row.index:
                base_name = self._safe_file_name(row[filename_column])
            elif self.markers and self.markers[0] in row.index:
                base_name = self._safe_file_name(row[self.markers[0]])
            else:
                base_name = f"row_{row_index + 1}"

            output_path = os.path.join(self.output_folder, f"Document_{base_name}.docx")
            doc.save(output_path)
            generated_files.append(output_path)

        return generated_files

    def get_all_tables(self):
        doc = Document(self.template)
        return doc.tables

    def find_libreoffice(self):

        candidatos = [
            "libreoffice",   # Linux / Mac (si está en PATH)
            "soffice",       # alternativo en PATH
            r"C:\Program Files\LibreOffice\program\soffice.exe",   # Windows
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # Mac
        ]
        for cmd in candidatos:
            if shutil.which(cmd) or os.path.isfile(cmd):
                return cmd
        return None


    def doc_to_pdf(self, doc_path, pdf_folder):

        soffice = self.find_libreoffice()
        if not soffice:
            raise EnvironmentError(
                "LibreOffice no encontrado. Instálalo desde https://www.libreoffice.org/"
            )

        result = subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", pdf_folder, doc_path],
            capture_output=True,
            text=True,
        )

        if result.returncode != 0:
            raise RuntimeError(f"Error al convertir {doc_path}:\n{result.stderr}")

        base_name = os.path.splitext(os.path.basename(doc_path))[0]
        return os.path.join(pdf_folder, base_name + ".pdf")


    def merge_pdfs(self, docx_path, pdf_final_path, order_file):

        pdf_folder_temp = os.path.join(docx_path, "_pdfs_temp")
        os.makedirs(pdf_folder_temp, exist_ok=True)

        writer = PdfWriter()
        errores = []

        for name, _ in order_file:
            file_name = name.replace(" ", "_").replace("/", "-")
            doc_path = os.path.join(docx_path, f"{file_name}.docx")

            if not os.path.exists(doc_path):
                print(f"  Not found: {os.path.basename(doc_path)}")
                continue

            try:
                ruta_pdf = self.doc_to_pdf(doc_path, pdf_folder_temp)
                reader = PdfReader(ruta_pdf)
                for page in reader.pages:
                    writer.add_page(page)
                print(f"{name}")
            except Exception as e:
                print(f"  {name}  →  {e}")
                errores.append((name, str(e)))

        with open(pdf_final_path, "wb") as f:
            writer.write(f)

        shutil.rmtree(pdf_folder_temp, ignore_errors=True)

        if errores:
            print(f"\n {len(errores)}")
            for name, err in errores:
                print(f"     • {name}: {err}")

        paginas_totales = sum(len(PdfReader(pdf_final_path).pages) for _ in [1])
        print(f"\n saved on '{pdf_final_path}' ({paginas_totales} pages)")
