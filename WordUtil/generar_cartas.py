"""
Generador de Cartas por Asesor
==============================
Reemplaza los marcadores NOMBRE_ASESOR y GENERO en una plantilla .docx,
genera un archivo Word por cada asesor y combina todos en un solo PDF.

Uso:
    python generar_cartas.py

Requisitos:
    pip install python-docx pypdf
    LibreOffice instalado (para la conversión a PDF)

    Windows:  https://www.libreoffice.org/download/download/
    Mac:      brew install --cask libreoffice
    Linux:    sudo apt install libreoffice
"""

from docx import Document
from pypdf import PdfWriter, PdfReader
import os
import glob
import subprocess
import sys
import shutil

# ─────────────────────────────────────────
# CONFIGURACIÓN — edita estos valores
# ─────────────────────────────────────────

PLANTILLA       = "carta_plantilla.docx"   # Ruta al archivo Word plantilla
CARPETA_SALIDA  = "cartas_generadas"        # Carpeta donde se guardarán los .docx
PDF_COMBINADO   = "todas_las_cartas.pdf"    # Nombre del PDF final combinado
MARCADOR_NOMBRE = "NOMBRE_ASESOR"          # Marcador de nombre en la carta
MARCADOR_GENERO = "GENERO"                 # Marcador de género en la carta

# Tuplas: (nombre completo, género "M" o "F")
ASESORES = [
    ("ING. LESLIE VIANNEY MARAÑON LIZARRAGA",           "F"),
    ("MC. IVETTE SELENE MARAÑON LIZARRAGA",              "F"),
    ("DRA. SHEILA SUSET MARAÑON LIZARRAGA",              "F"),
    ("MC. FERNANDO RAFAEL PRADO VALENZUELA",             "M"),
    ("MC. ALEJANDRA GUADALUPE RODRIGUEZ DOMINGUEZ",      "F"),
    ("DR. FRANCISCO HUMBERTO VALDEZ SANDOVAL",           "M"),
    ("DRA. KARLA VIANNEY RUELAS AGUILAR",                "F"),
    ("DRA. NANCY CECILIA MITCHEL RIVERA",                "F"),
    ("LIC. ALEJANDRA SOTO CHAVEZ",                       "F"),
    ("DR. JOSE CRISOFORO CARRAZCO ESCALANTE",            "M"),
    ("LIC. LUIS FERNANDO LOPEZ PARRA",                   "M"),
    ("DR. FRANCISCO GUILLERMO SALCIDO VEGA",             "M"),
    ("DRA. CRISTIAN ARMENTILLA GALAVIZ",                 "F"),
    ("DRA. MARISOL ARMENTILLA GALAVIZ",                  "F"),
    ("LIC. WENDY LIZBETH QUIÑONEZ ARGUELLO",             "F"),
    ("DRA. MARIA GUADALUPE GAMEZ MEDINA",                "F"),
    ("DRA. YOBANA DAMARIZ CORTEZ PARRA",                 "F"),
    ("MC. SAMUEL MORENO MEZA",                           "M"),
    ("MBA. DAVID RENE YEPIS PEÑUELAS",                   "M"),
    ("LIC. JOSE ALBERTO VELAZQUEZ CASTRO",               "M"),
    ("LIC. MARIA GUADALUPE VEGA NUÑEZ",                  "F"),
    ("DR. ANGEL ESTEBAN MANZANAREZ SALAZAR",             "M"),
    ("MC. MARIA DE JESUS HAMASAKI GALVEZ",               "F"),
    ("DRA. FLOR ALMA WONG MONTIEL",                      "F"),
    ("DRA. MIRIAM PEREZ BARRAZA",                        "F"),
    ("DR. GUSTAVO FABIAN PEREZ",                         "M"),
    ("DR. JORGE EUGENIO DE JESUS MORA TORDECILLAS",      "M"),
]

# ─────────────────────────────────────────


def genero_a_titulo(genero):
    return "Evaluadora" if genero == "F" else "Evaluador"


def reemplazar_en_parrafo(parrafo, reemplazos):
    texto_completo = "".join(run.text for run in parrafo.runs)
    if not any(m in texto_completo for m in reemplazos):
        return
    texto_nuevo = texto_completo
    for marcador, valor in reemplazos.items():
        texto_nuevo = texto_nuevo.replace(marcador, valor)
    if parrafo.runs:
        parrafo.runs[0].text = texto_nuevo
        for run in parrafo.runs[1:]:
            run.text = ""


def reemplazar_en_doc(doc, reemplazos):
    for parrafo in doc.paragraphs:
        reemplazar_en_parrafo(parrafo, reemplazos)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for parrafo in celda.paragraphs:
                    reemplazar_en_parrafo(parrafo, reemplazos)
    for seccion in doc.sections:
        for contenedor in [
            seccion.header, seccion.first_page_header, seccion.even_page_header,
            seccion.footer, seccion.first_page_footer, seccion.even_page_footer,
        ]:
            if contenedor is None:
                continue
            for parrafo in contenedor.paragraphs:
                reemplazar_en_parrafo(parrafo, reemplazos)
            for tabla in contenedor.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for parrafo in celda.paragraphs:
                            reemplazar_en_parrafo(parrafo, reemplazos)


def generar_carta(plantilla_path, nombre, genero, carpeta_salida):
    doc = Document(plantilla_path)
    reemplazos = {
        MARCADOR_NOMBRE: nombre,
        MARCADOR_GENERO: genero_a_titulo(genero),
    }
    reemplazar_en_doc(doc, reemplazos)
    nombre_archivo = nombre.replace(" ", "_").replace("/", "-")
    ruta_salida = os.path.join(carpeta_salida, f"Carta_{nombre_archivo}.docx")
    doc.save(ruta_salida)
    return ruta_salida


# ─────────────────────────────────────────
# Conversión a PDF y combinación
# ─────────────────────────────────────────

def encontrar_libreoffice():
    """Busca el ejecutable de LibreOffice en las rutas comunes."""
    candidatos = [
        "libreoffice",   # Linux / Mac (si está en PATH)
        "soffice",       # alternativo en PATH
        r"C:\Program Files\LibreOffice\program\soffice.exe",   # Windows
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # Mac
    ]
    for cmd in candidatos:
        if shutil.which(cmd) or os.path.isfile(cmd):
            return cmd
    return None


def docx_a_pdf(ruta_docx, carpeta_pdf):
    """Convierte un .docx a PDF usando LibreOffice y retorna la ruta del PDF."""
    soffice = encontrar_libreoffice()
    if not soffice:
        raise EnvironmentError(
            "LibreOffice no encontrado. Instálalo desde https://www.libreoffice.org/"
        )

    resultado = subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", carpeta_pdf, ruta_docx],
        capture_output=True,
        text=True,
    )

    if resultado.returncode != 0:
        raise RuntimeError(f"Error al convertir {ruta_docx}:\n{resultado.stderr}")

    nombre_base = os.path.splitext(os.path.basename(ruta_docx))[0]
    return os.path.join(carpeta_pdf, nombre_base + ".pdf")


def combinar_pdfs(carpeta_docx, ruta_pdf_final, orden_asesores):
    """
    Convierte todos los .docx a PDF (en el mismo orden que la lista de asesores)
    y los combina en un único archivo PDF.
    """
    print("\n📑 Combinando cartas en un solo PDF...")
    carpeta_pdf_temp = os.path.join(carpeta_docx, "_pdfs_temp")
    os.makedirs(carpeta_pdf_temp, exist_ok=True)

    writer = PdfWriter()
    errores = []

    for nombre, _ in orden_asesores:
        nombre_archivo = nombre.replace(" ", "_").replace("/", "-")
        ruta_docx = os.path.join(carpeta_docx, f"Carta_{nombre_archivo}.docx")

        if not os.path.exists(ruta_docx):
            print(f"  ⚠️  No encontrado, omitido: {os.path.basename(ruta_docx)}")
            continue

        try:
            ruta_pdf = docx_a_pdf(ruta_docx, carpeta_pdf_temp)
            reader = PdfReader(ruta_pdf)
            for page in reader.pages:
                writer.add_page(page)
            print(f"  ✅  {nombre}")
        except Exception as e:
            print(f"  ❌  {nombre}  →  {e}")
            errores.append((nombre, str(e)))

    with open(ruta_pdf_final, "wb") as f:
        writer.write(f)

    # Limpiar PDFs temporales
    shutil.rmtree(carpeta_pdf_temp, ignore_errors=True)

    if errores:
        print(f"\n  ⚠️  {len(errores)} carta(s) no se pudieron convertir:")
        for nombre, err in errores:
            print(f"     • {nombre}: {err}")

    paginas_totales = sum(len(PdfReader(ruta_pdf_final).pages) for _ in [1])
    print(f"\n✔ PDF combinado guardado: '{ruta_pdf_final}' ({paginas_totales} páginas)")


# ─────────────────────────────────────────
# Main
# ─────────────────────────────────────────

def main():
    if not os.path.exists(PLANTILLA):
        print(f"❌ No se encontró la plantilla: '{PLANTILLA}'")
        print("   Asegúrate de que el archivo esté en la misma carpeta que este script.")
        return

    os.makedirs(CARPETA_SALIDA, exist_ok=True)

    print(f"📄 Plantilla:      {PLANTILLA}")
    print(f"📁 Carpeta salida: {CARPETA_SALIDA}")
    print(f"📋 PDF combinado:  {PDF_COMBINADO}")
    print(f"👥 Asesores:       {len(ASESORES)}\n")
    print("Generando cartas Word...")
    print("─" * 50)

    errores_docx = []
    for nombre, genero in ASESORES:
        try:
            generar_carta(PLANTILLA, nombre, genero, CARPETA_SALIDA)
            print(f"  ✅  [{genero_a_titulo(genero):<12}]  {nombre}")
        except Exception as e:
            print(f"  ❌  {nombre}  →  {e}")
            errores_docx.append((nombre, str(e)))

    print("─" * 50)
    generadas = len(ASESORES) - len(errores_docx)
    print(f"\n✔ {generadas} cartas Word generadas en '{CARPETA_SALIDA}/'")

    # Combinar en PDF
    combinar_pdfs(CARPETA_SALIDA, PDF_COMBINADO, ASESORES)


if __name__ == "__main__":
    main()
